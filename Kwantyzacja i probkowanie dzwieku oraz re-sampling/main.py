from docx import Document
from docx.shared import Inches
import numpy as np
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
import scipy.fftpack
import soundfile as sf

def kwantyzacja(data, bits):
    if (bits < 2 or bits >32) :
        raise ValueError("Liczba bitów musi być z zakresu od 2 do 32.")

    output = data.astype(np.float32)
    min = 0
    max = 0

    if np.issubdtype(data.dtype, np.floating):
        min = -1
        max= 1
    else:
        min = np.iinfo(data.dtype).min
        max   = np.iinfo(data.dtype).max

    d = 2**bits-1

    output = (output-min)/(max-min)
    output = np.round(output*d)/d
    output = ((output*(max-min))+min).astype(data.dtype)
    return output

def decymacja(signal, N):
    return signal[::N].copy()

def interpolacja(data, freq, new_freq, interp_type="linear"):
    time = data.shape[0]/freq
    t  = np.linspace(0, time, len(data))
    t1 = np.linspace(0, time, int(time * new_freq))

    interp = interp1d(t, data, interp_type, fill_value="extrapolate")
    return interp(t1).astype(data.dtype)

def plotAudio(Signal, Fs, TimeMargin=[0, 0.02], f_size=2 ** 8, Axs=None, title=None): 
    if Axs is None:
        _, Axs = plt.subplots(2, 1, figsize=(10, 7))
    Axs[0].plot(np.arange(0, len(Signal)) / Fs, Signal)
    Axs[0].set_xlim(TimeMargin[0], TimeMargin[1])
    yf = scipy.fftpack.fft(Signal, f_size)
    xf = np.arange(0, Fs / 2, Fs / f_size)
    spec = 20 * np.log10(np.abs(yf[:f_size // 2]))
    Axs[1].plot(xf, spec)
    if title is not None:
        plt.suptitle(title)  
    return xf[np.argmax(spec)], np.max(spec)

document = Document()

files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav','sin_combined.wav']
files2 = ['sing_high1.wav','sing_medium1.wav','sing_low1.wav']

for file in files:
    signal, samplerate = sf.read(file)
    
    for bits in [4, 8, 16, 24]:
        modified_signal = kwantyzacja(signal, bits)
        fig, ax = plt.subplots(2, 1, figsize=(10, 7))
        plotAudio(modified_signal, samplerate, title=f'{file} - {bits}-bit', Axs=ax)
        plt.close(fig)  
        img_path = f"temp_plot_{file}_{bits}bit.png"
        fig.savefig(img_path)
        document.add_picture(img_path, width=Inches(6))
        plt.close()

    for decimation_factor in [2, 4, 6, 10, 24]:
        decimated_signal = decymacja(signal, decimation_factor)
        fig, ax = plt.subplots(2, 1, figsize=(10, 7))
        plotAudio(decimated_signal, samplerate // decimation_factor, title=f'{file} - decymacja {decimation_factor}', Axs=ax)
        plt.close(fig)
        img_path = f"temp_plot_{file}_decymacja{decimation_factor}.png"
        fig.savefig(img_path)
        document.add_picture(img_path, width=Inches(6))
        plt.close()

    for new_freq in [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]:
        for interp_type in ["linear", "cubic"]:
            interpolated_signal = interpolacja(signal, samplerate, new_freq, interp_type)
            fig, ax = plt.subplots(2, 1, figsize=(10, 7))
            plotAudio(interpolated_signal, new_freq, title=f'{file} - interpolacja {new_freq} Hz - {interp_type}', Axs=ax)
            plt.close(fig)
            img_path = f"temp_plot_{file}_interpolacja{new_freq}_{interp_type}.png"
            fig.savefig(img_path)
            document.add_picture(img_path, width=Inches(6))
            plt.close()

    for bits in [4, 8]:
        modified_signal = kwantyzacja(signal, bits)
        sf.write(f"{file[:-4]}_{bits}-bit.wav", modified_signal, samplerate)

    for decimation_factor in [4, 6, 10, 24]:
        decimated_signal = decymacja(signal, decimation_factor)
        sf.write(f"{file[:-4]}_decymacja{decimation_factor}.wav", decimated_signal, samplerate // decimation_factor)

    for new_freq in [4000, 8000, 11999, 16000, 16953]:
        for interp_type in ["linear", "cubic"]:
            interpolated_signal = interpolacja(signal, samplerate, new_freq, interp_type)
            sf.write(f"{file[:-4]}_interpolacja{new_freq}_{interp_type}.wav", interpolated_signal, new_freq)

document.save('raport.docx')

