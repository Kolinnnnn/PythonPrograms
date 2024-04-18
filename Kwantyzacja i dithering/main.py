import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from io import BytesIO

def imgToFloat(img):
    if np.issubdtype(img.dtype, np.unsignedinteger):
        return img / 255.0
    elif np.issubdtype(img.dtype, np.floating):
        return img
    else:
        raise ValueError("Input img type not supported")

GS1 = imgToFloat(plt.imread('GS_0001.tif'))
GS2 = imgToFloat(plt.imread('GS_0002.png'))
GS3 = imgToFloat(plt.imread('GS_0003.png'))
SMALL1 = imgToFloat(plt.imread('SMALL_0001.tif'))
SMALL2 = imgToFloat(plt.imread('SMALL_0002.png'))
SMALL3 = imgToFloat(plt.imread('SMALL_0003.png'))
SMALL4 = imgToFloat(plt.imread('SMALL_0004.jpg'))
SMALL5 = imgToFloat(plt.imread('SMALL_0005.jpg'))
SMALL6 = imgToFloat(plt.imread('SMALL_0006.jpg'))
SMALL7 = imgToFloat(plt.imread('SMALL_0007.jpg'))
SMALL8 = imgToFloat(plt.imread('SMALL_0008.jpg'))
SMALL9 = imgToFloat(plt.imread('SMALL_0009.jpg'))
SMALL10 = imgToFloat(plt.imread('SMALL_0010.jpg'))

GS1_gray = GS1[:, :, 0]
GS2_gray = GS2[:, :, 0]
GS3_gray = GS3[:, :, 0]

Graypaleta1bit = np.linspace(0, 1, 2).reshape((2, 1))
Graypaleta2bit = np.linspace(0, 1, 4).reshape((4, 1))
Graypaleta4bit = np.linspace(0, 1, 8).reshape((8, 1))

paleta8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],
])

paleta16 =  np.array([
        [0.0, 0.0, 0.0,], 
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,], 
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,], 
        [1.0, 1.0, 0.0,]
])

def colorFit(color, paleta):
    x = np.linalg.norm(paleta - color, axis=1)
    return paleta[np.argmin(x)]

def kwant_colorFit(img,paleta):
    out_img = img.copy()
    for w in range(0, img.shape[0]):
            for k in range(0, img.shape[1]):
                    out_img[w,k]=colorFit(img[w,k],paleta)
    return out_img

def random(img):
    out_img = img.copy()
    r = np.random.rand(img.shape[0], img.shape[1])
    for row in range(0, img.shape[0]):
        for col in range(0, img.shape[1]):
            out_img[row, col] = img[row, col] > r[row, col]
    return out_img

def organized(img, paleta):
    out_img = img.copy()
    m2 = np.matrix('0,8,2,10;'
                   '12,4,14,6;'
                   '3,11,1,9;'
                   '15,7,13,5')
    Mpre = 1 / 16 * (m2 + 1) - 0.5
    for row in range(0, img.shape[0]):
        for col in range(0, img.shape[1]):
            out_img[row, col] = colorFit(out_img[row, col] + Mpre[row % 4, col % 4], paleta)
    return out_img

def floydSteinberg(img, paleta):
    out_img = img.copy()
    for row in range(1, img.shape[0]-1):
        for col in range(1, img.shape[1]-1):
            oldpixel = out_img[row, col].copy()
            newpixel = colorFit(oldpixel, paleta)
            out_img[row, col] = newpixel
            quanterror = oldpixel - newpixel
            out_img[row + 1, col    ] = np.clip(out_img[row + 1, col    ] + quanterror * 7 / 16, 0, 1)
            out_img[row - 1, col + 1] = np.clip(out_img[row - 1, col + 1] + quanterror * 3 / 16, 0, 1)
            out_img[row    , col + 1] = np.clip(out_img[row    , col + 1] + quanterror * 5 / 16, 0, 1)
            out_img[row + 1, col + 1] = np.clip(out_img[row + 1, col + 1] + quanterror * 1 / 16, 0, 1)
    return out_img

def show_and_save(img, palette, file_name, bit_depth, document):
    if len(palette) == 2 or 4 or 8:
        cmap = plt.cm.gray

    plt.figure(figsize=(10, 5))

    plt.subplot(2, 3, 1)
    plt.title('Oryginal')
    plt.imshow(img, cmap=cmap)

    plt.subplot(2, 3, 2)
    plt.title('Kwantyzacja')
    plt.imshow(kwant_colorFit(img, palette), cmap=cmap)

    plt.subplot(2, 3, 3)
    plt.title('Dithering zorganizowany')
    plt.imshow(organized(img, palette), cmap=cmap)

    plt.subplot(2, 3, 4)
    plt.title('Dithering Floyd-Steinberg')
    plt.imshow(floydSteinberg(img, palette), cmap=cmap)

    if len(palette) == 2:
        plt.subplot(2, 3, 6)
        plt.title('Dithering losowy')
        plt.imshow(random(img), cmap=cmap)

    plt.tight_layout()

    memfile = BytesIO()
    plt.savefig(memfile)
    memfile.seek(0)

    document.add_paragraph('FILE: {}'.format(file_name))
    document.add_paragraph('Dithering {}-bit'.format(bit_depth))

    document.add_picture(memfile, width=Inches(6))
    
    memfile.close()

    document.add_page_break()

document = Document()

show_and_save(GS1, Graypaleta1bit, 'GS_0001.tif', 1, document)
show_and_save(GS1, Graypaleta2bit, 'GS_0001.tif', 2, document)
show_and_save(GS1, Graypaleta4bit, 'GS_0001.tif', 4, document)
show_and_save(GS2_gray, Graypaleta1bit, 'GS_0002.png', 1, document)
show_and_save(GS2_gray, Graypaleta2bit, 'GS_0002.png', 2, document)
show_and_save(GS2_gray, Graypaleta4bit, 'GS_0002.png', 4, document)
show_and_save(GS3, Graypaleta1bit, 'GS_0003.png', 1, document)
show_and_save(GS3, Graypaleta2bit, 'GS_0003.png', 2, document)
show_and_save(GS3, Graypaleta4bit, 'GS_0003.png', 4, document)
show_and_save(SMALL7, paleta8, 'SMALL_0007.jpg', 8, document)
show_and_save(SMALL7, paleta16, 'SMALL_0007.jpg', 16, document)
show_and_save(SMALL4, paleta8, 'SMALL_0004.jpg', 8, document)
show_and_save(SMALL4, paleta16, 'SMALL_0004.jpg', 16, document)
show_and_save(SMALL9, paleta8, 'SMALL_0009.jpg', 8, document)
show_and_save(SMALL9, paleta16, 'SMALL_0009.jpg', 16, document)
show_and_save(SMALL6, paleta8, 'SMALL_0006.jpg', 8, document)
show_and_save(SMALL6, paleta16, 'SMALL_0006.jpg', 16, document)

document.save('raport.docx')
